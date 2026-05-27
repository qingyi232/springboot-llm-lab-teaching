# -*- coding: utf-8 -*-
"""生成毕业答辩常见问题与回答 docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = r"F:\26毕设2\基于SpringBoot和大模型的高校实验教学支持系统的设计与实现\毕业答辩常见问题与回答.docx"


def set_run_font(run, name="宋体", size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_title(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level > 1 else 18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    sizes = {0: 22, 1: 16, 2: 14, 3: 12}
    set_run_font(run, "黑体", sizes.get(level, 12), bold=True)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, "宋体", 12)


def add_qa(doc, q, a, tips=None):
    pq = doc.add_paragraph()
    pq.paragraph_format.space_before = Pt(8)
    rq = pq.add_run("问：" + q)
    set_run_font(rq, "宋体", 12, bold=True)

    pa = doc.add_paragraph()
    pa.paragraph_format.first_line_indent = Cm(0.74)
    pa.paragraph_format.line_spacing = 1.5
    ra = pa.add_run("答：" + a)
    set_run_font(ra, "宋体", 12)

    if tips:
        pt = doc.add_paragraph()
        pt.paragraph_format.first_line_indent = Cm(0.74)
        rt = pt.add_run("【答辩技巧】" + tips)
        set_run_font(rt, "宋体", 11)
        rt.font.color.rgb = RGBColor(0x80, 0x40, 0x00)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2.5)

    add_title(doc, "毕业答辩常见问题与回答", 0)
    add_title(doc, "基于SpringBoot和大模型的高校实验教学支持系统的设计与实现", 1)
    add_body(doc, "说明：本文档结合项目源码（backend/frontend）、数据库脚本 init.sql、README 及论文摘要整理，供答辩前背诵与模拟问答使用。演示账号密码均为 123456。")

    # 一、必背要点
    add_title(doc, "一、答辩前必背的 8 个核心点（30 秒版）", 1)
    points = [
        "选题背景：高校实验教学资源紧张、预约冲突、反馈滞后，需要信息化+智能化平台。",
        "技术栈：后端 Spring Boot 2.7 + MyBatis-Plus + MySQL；前端 Vue3 + Element Plus；AI 集成 DeepSeek，SSE 流式输出。",
        "架构：前后端分离，RESTful API，context-path 为 /api，端口 8080（后端）、5173（前端）。",
        "四类角色：管理员、教师、实验室管理员、学生，Token 表鉴权 + 路由按角色分流。",
        "核心功能：实验室预约与审核、课程资料/视频、作业提交与成绩、论坛、公告、收藏、AI 智能问答。",
        "数据库：约 15+ 张业务表（student/teacher/lab_info/reservation/homework/forum 等），见 init.sql。",
        "创新点（按论文表述）：大模型辅助教学问答；多角色协同；实验资源在线化管理。（若老师追问实现细节，见下文 AI 模块说明）",
        "测试结论：功能测试、黑盒测试验证主要流程可用，系统运行稳定。",
    ]
    for i, t in enumerate(points, 1):
        add_body(doc, f"{i}. {t}")

    # 二、选题与意义
    add_title(doc, "二、选题背景、研究意义与国内外现状", 1)
    add_qa(doc,
        "你为什么选这个题目？有什么实际意义？",
        "实验教学是高校人才培养的重要环节。随着办学规模扩大，实验室资源紧张、预约靠人工、资料分散、师生互动不及时等问题突出。本课题设计并实现一套实验教学支持系统，把预约、资料、作业、交流集中到线上，并引入大模型提供智能问答，提高管理效率和教学支持能力，对高校实验教学信息化有实际应用价值。",
        "先讲痛点（3 句）→ 再讲你的系统解决什么 → 最后一句意义，不要一上来背技术名词。")

    add_qa(doc,
        "你的系统和普通教务系统/实验室管理系统有什么区别？",
        "普通教务系统侧重课表、成绩等行政流程；本系统聚焦实验教学场景：实验室资源展示与预约审核、实验课程资料与视频、作业布置提交批改、师生论坛交流，并集成 DeepSeek 大模型做实验与课程相关的智能辅导，更贴近实验课教与学的全过程。",
        "突出「实验教学专属」+「AI 辅助」，避免说「功能更多」这种空话。")

    add_qa(doc,
        "论文摘要里提到智能调度、个性化推荐，系统里怎么实现的？",
        "论文从研究目标上阐述了利用大模型分析使用数据、优化预约与教学决策的方向。当前已实现的核心 AI 能力是：基于 DeepSeek API 的智能问答（含 SSE 流式回复），系统提示词限定为实验教学助手，可解答课程、实验操作、预约与作业等问题；未配置 API Key 时提供离线模拟回答保证演示。智能调度与推荐可作为后续扩展，答辩时如实说明「已实现问答辅导，数据分析与推荐为设计与展望」。",
        "这是高频刁钻题，必须诚实：已实现什么、论文展望什么，分开说，老师反而认可。")

    # 三、技术架构
    add_title(doc, "三、技术架构与选型", 1)
    add_qa(doc,
        "请介绍系统的总体架构。",
        "采用 B/S 架构、前后端分离。前端 Vue3 单页应用通过 Axios 调用后端 REST 接口；后端 Spring Boot 提供 Controller-Service-Mapper 分层，MyBatis-Plus 访问 MySQL；文件上传存本地 uploads 目录并通过静态资源映射访问；AI 模块独立为 DeepSeekController，对接 DeepSeek Chat Completions API。",
        "建议画三层图：表现层 Vue → 业务层 Spring Boot → 数据层 MySQL。")

    add_qa(doc,
        "为什么选 Spring Boot？",
        "Spring Boot 开箱即用、生态成熟，适合快速构建 REST 服务；内嵌 Tomcat 部署简单；与 MyBatis-Plus、Druid 数据源等集成方便，符合本科毕设开发周期与可维护性要求。",
        None)

    add_qa(doc,
        "为什么前端用 Vue3 而不是 React 或小程序？",
        "Vue3 组合式 API 学习曲线平缓，与 Element Plus 后台/前台组件库配合成熟，适合管理端+学生端双端界面；Vite 构建速度快。本系统面向 PC 浏览器使用场景，故采用 Web 前端而非小程序。",
        None)

    add_qa(doc,
        "为什么用 MySQL？有没有考虑 Redis？",
        "业务数据以结构化为主（用户、预约、作业、论坛等），MySQL 关系型存储合适，init.sql 定义完整表结构。application.yml 中配置了 Redis，可用于后续缓存热点数据或会话；当前登录鉴权主要使用数据库 token 表存储令牌与过期时间。",
        "如实说 Redis 配置存在、业务主路径用 MySQL+Token 表。")

    add_qa(doc,
        "前后端怎么交互？跨域怎么处理？",
        "前端开发服务器通过 Vite 代理或直连 http://localhost:8080/api；请求 Header 携带 Token（登录后写入 localStorage）。后端 AuthInterceptor 校验 Token 有效性；登录、注册、部分列表接口在 WebMvcConfig 中排除拦截。AI 聊天使用 EventSource 订阅 /api/ai/chat/sse 实现流式输出。",
        None)

    # 四、功能模块
    add_title(doc, "四、功能模块与业务流程", 1)
    add_qa(doc,
        "系统有哪些功能模块？",
        "①用户与权限：四类角色注册登录；②实验室：类型、信息、预约申请、审核、使用记录；③教学资源：课程类型、资料、视频；④作业：发布、提交、成绩评语；⑤交流：论坛发帖回帖；⑥公告；⑦收藏；⑧后台管理；⑨AI 智能问答。",
        "按角色讲更清晰：学生用什么、教师用什么、实验室管理员用什么。")

    add_qa(doc,
        "实验室预约的完整流程是什么？",
        "学生在实验室列表查看详情 → 填写预约时间、人数、说明 → 提交后 reservation 表状态为「待审核」→ 管理员或实验室管理员在后台审核（通过/回复）→ 通过后学生可在「我的预约」查看；使用后登记使用记录 usage_record。",
        "强调「待审核」体现管理规范，可对比论文中解决预约冲突。")

    add_qa(doc,
        "作业模块怎么实现的？",
        "教师发布 homework（标题、内容、附件、截止时间、班级）；学生在截止前 homework_submit 提交内容与附件；教师或管理员在 homework_score 录入分数与 pingyu_content；学生端「我的成绩」查看。",
        None)

    add_qa(doc,
        "论坛模块有什么作用？",
        "forum 表支持发帖与 parentid 回复，实现师生异步交流，可讨论实验问题、课程疑问，管理员可后台管理帖子，形成教学社区氛围。",
        None)

    # 五、数据库
    add_title(doc, "五、数据库设计", 1)
    add_qa(doc,
        "数据库有哪些核心表？表之间什么关系？",
        "核心表包括：users（管理员）、student、teacher、lab_admin、token；lab_type、lab_info；reservation、usage_record；course_type、course_material、video；homework、homework_submit、homework_score；forum、notice、collection 等。关系上：预约/使用记录关联实验室与学生信息（冗余 lab_id、xuehao 等字段便于查询）；作业提交关联 homework_id；成绩关联 submit_id。",
        "准备画 E-R 图：用户-预约-实验室、教师-作业-提交-成绩。")

    add_qa(doc,
        "为什么预约表要冗余 lab_name、xingming 等字段？",
        "属于空间换时间的设计，列表与审核页可直接展示，减少多表 JOIN，提高后台查询效率；缺点是实验室改名需同步更新历史记录，毕设规模可接受。",
        None)

    # 六、AI
    add_title(doc, "六、大模型（DeepSeek）集成", 1)
    add_qa(doc,
        "大模型在系统中怎么用的？",
        "通过 DeepSeekController 调用 DeepSeek Chat API。配置项：api-key、api-url、model（deepseek-chat）。提供两种接口：POST /ai/chat 一次性返回；GET /ai/chat/sse 使用 SSE 流式推送，前端 AiChat.vue 用 EventSource 逐字显示。System 提示词将模型限定为「高校实验教学 AI 助手」。",
        "答辩可现场演示 AI 页面；提前确认 API Key 有效或说明演示模式。")

    add_qa(doc,
        "什么是 SSE？为什么用流式输出？",
        "SSE（Server-Sent Events）是服务器向浏览器单向推送事件的技术。大模型生成较慢，流式输出可让用户实时看到回答，体验更好；后端将完整回复分块写入 data: {\"content\":\"...\"} 格式，结束时发送 data: [DONE]。",
        None)

    add_qa(doc,
        "没有 API Key 时系统还能用吗？",
        "可以。代码中判断 apiKey 为空或以 sk-your 开头时走 generateMockReply，按关键词（数据结构、实验预约、作业成绩等）返回预设辅导内容，保证答辩演示不中断。",
        "主动说明，体现你考虑容灾。")

    add_qa(doc,
        "大模型会不会胡编答案？你怎么控制？",
        "通过 System 角色提示约束回答范围；temperature 设为 0.7 平衡多样与稳定；前端展示为「辅助参考」而非标准答案。作业仍由教师批改，AI 不替代成绩评定。可补充：敏感操作需登录，API Key 放服务端配置不暴露前端。",
        "老师关心伦理与准确性，要答「辅助」定位。")

    # 七、安全
    add_title(doc, "七、安全、权限与性能", 1)
    add_qa(doc,
        "系统怎么实现登录与权限控制？",
        "各角色独立登录接口，验证通过后 TokenService 生成 token 写入 token 表（含 userid、role、tablename、过期时间）。后续请求 Header 带 Token，AuthInterceptor 校验存在且未过期，并将 userId、role 放入 request。前端路由 meta.role 区分管理员/教师/实验室管理员入口。",
        None)

    add_qa(doc,
        "密码怎么存储？",
        "毕设演示数据为明文或简单存储（init.sql 演示账号）；正式环境应使用 BCrypt 等哈希加盐存储，答辩可答「演示环境简化，生产需加密」。",
        "诚实回答，提出改进方案加分。")

    add_qa(doc,
        "文件上传有什么限制？",
        "application.yml 配置 multipart 最大 100MB；FileController 保存到 file.upload-path（默认 ./uploads/），通过 /uploads/** 静态访问。",
        None)

    # 八、测试与创新
    add_title(doc, "八、测试、难点、创新点与不足", 1)
    add_qa(doc,
        "你做了哪些测试？",
        "按论文：功能测试覆盖登录、预约审核、资料浏览、作业提交、论坛、AI 问答等主流程；黑盒测试验证输入边界与异常提示；前后端联调测试接口返回码与页面展示。",
        "准备 2～3 个测试用例表截图。")

    add_qa(doc,
        "开发中遇到的最大难点是什么？",
        "可参考：①前后端 Token 鉴权与拦截器路径排除配置；②AI SSE 流式数据格式与前端拼接显示；③预约时间段与状态流转的业务一致性；④多角色菜单与路由权限分离。选 1 个讲清「问题-排查-解决」。",
        "用 STAR：情境-任务-行动-结果。")

    add_qa(doc,
        "系统的创新点是什么？",
        "①将实验教学预约、资源、作业、交流整合于统一平台；②多角色协同（管理员、实验室管理员、教师、学生）；③集成 DeepSeek 大模型提供实验教学场景智能问答与流式交互；④前后端分离便于扩展维护。",
        None)

    add_qa(doc,
        "系统有什么不足和后续改进？",
        "不足：①大模型目前以问答为主，智能推荐与资源调度尚未完全落地；②Redis 未深度用于缓存；③密码存储需加强加密；④移动端适配可优化。改进：对接校内统一认证、增加预约冲突自动检测、RAG 结合课程资料库增强问答准确性、引入数据统计大屏。",
        "主动说不足显得成熟，不要辩称「没有不足」。")

    # 九、刁钻题
    add_title(doc, "九、答辩高频「刁钻」问题", 1)
    add_qa(doc,
        "代码是你自己写的吗？",
        "核心业务流程、接口设计与 AI 对接部分本人完成开发与调试；使用了 Spring Boot、Vue、Element Plus 等成熟框架和 MyBatis-Plus 简化 CRUD。能现场演示登录、预约、发帖、AI 问答，并说明 Controller/表结构对应关系。",
        "自信、具体，可指到 DeepSeekController、init.sql。")

    add_qa(doc,
        "如果并发很多人同时预约同一实验室怎么办？",
        "当前通过预约审核与状态字段控制；高并发场景可在提交时加数据库行锁或 Redis 分布式锁，并校验时间段容量与 rongna（容纳人数）。毕设阶段以流程正确为主，并发优化可作为展望。",
        None)

    add_qa(doc,
        "为什么用 DeepSeek 不用 ChatGPT？",
        "DeepSeek 提供 OpenAI 兼容的 Chat Completions 接口，国内访问与成本更适合毕设；Spring RestTemplate 即可对接，便于 SSE 封装与论文「国产大模型」表述一致。",
        None)

    add_qa(doc,
        "你的系统和论文描述不一致怎么办？",
        "论文从研究目标阐述智能化分析与推荐；系统已实现平台化管理与 AI 问答核心功能。对未完全实现部分如实说明为后续工作，不夸大。",
        "态度诚恳最重要。")

    # 十、答辩准备
    add_title(doc, "十、答辩当天准备清单", 1)
    checklist = [
        "演示环境：MySQL 已导入 init.sql；后端 8080、前端 5173 能启动；DeepSeek Key 有效或接受模拟模式。",
        "账号准备：admin / 2021001 / T2024001 / G2024001，密码 123456。",
        "演示顺序建议：登录学生 → 浏览实验室 → 提交预约 → 管理员审核 → 作业/论坛 → AI 问答（SSE 流式）。",
        "材料：论文纸质版、U 盘备份项目与数据库脚本、架构图/功能结构图打印 1 页。",
        "陈述控制 8～10 分钟：背景 1min → 功能演示 4min → 技术+AI 2min → 测试总结 1min。",
        "穿着得体，语速适中；不会的问题答「该点后续改进，感谢老师指正」。",
        "特别注意：老师常追问 AI 实现细节、表结构设计、你个人负责模块，提前熟读本文档第二节与第六节。",
    ]
    for t in checklist:
        add_body(doc, "• " + t)

    add_title(doc, "附录：演示账号与技术参数速查", 2)
    add_body(doc, "演示账号：管理员 admin；学生 2021001；教师 T2024001；实验室管理员 G2024001；密码均为 123456。")
    add_body(doc, "后端：Spring Boot 2.7.18，端口 8080，context-path /api，数据库 lab_teaching。")
    add_body(doc, "前端：Vue3 + Vite + Element Plus，端口 5173。")
    add_body(doc, "AI：DeepSeek API，接口 /api/ai/chat 与 /api/ai/chat/sse。")

    doc.save(OUT)
    print("已生成:", OUT)


if __name__ == "__main__":
    build()
